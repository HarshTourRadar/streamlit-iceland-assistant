import asyncio
import json
from pathlib import Path

import requests
from docx import Document
from docx.shared import Inches
from PIL import Image


def image_to_jpg(image_path):
    path = Path(image_path)
    valid_images_extensions = {
        ".jpg",
        ".png",
        ".jfif",
        ".exif",
        ".gif",
        ".tiff",
        ".bmp",
    }
    if path.suffix not in valid_images_extensions:
        jpg_image_path = f"{path.parent / path.stem}_result.jpg"
        Image.open(image_path).convert("RGB").save(jpg_image_path)
        return jpg_image_path
    return image_path


async def download_images(url: str, document):
    response = requests.get(url)
    if response.status_code == 200:
        # Proceed to save the image
        with open("image.jpeg", "wb") as f:
            f.write(response.content)
        document.add_picture(image_to_jpg("image.jpeg"), width=Inches(2.25))
    else:
        print("Failed to download image. Status code:", response.status_code)
    return


async def read_json():
    with open("./iceland_adventures.json", "r") as f:
        data = json.load(f)
        total_tours = len(data)
        docs_counter = 0
        tours_per_doc = 10
        for i in range(0, total_tours, tours_per_doc):
            document = Document()
            tours_chunk = data[i : i + tours_per_doc]
            for index, tour in enumerate(tours_chunk):
                document.add_heading(
                    f"Tour ID: {tour['tour_id']}\nTour Name: {tour['tour_name']}",
                    0,
                )
                destinations_list = [
                    city["city_name"] for city in tour["destinations"]["cities"]
                ]
                document.add_paragraph(f"Tour destinations: {destinations_list}")
                document.add_paragraph("Tour destinations images:")
                images = tour["images"]
                for url in images:
                    await download_images(url["url"], document)
                    print(f"Image downloaded: ({i+index+1}/{total_tours})")
                document.add_paragraph("\n")
                print("\n")
            docs_counter += 1
            document.save(f"tour_images_part_{docs_counter}.docx")


async def main():
    await read_json()


if __name__ == "__main__":
    asyncio.run(main())
